"""
DOCX Resume Parser — Text extraction from Word documents.
Handles paragraphs, tables, headers, and edge cases.
"""
from docx import Document
import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> dict:
    """
    Extract text from a DOCX file.
    
    Returns:
        dict with keys:
            - text: extracted full text
            - sections: list of section texts
            - success: bool
            - error: error message if any
    """
    result = {
        "text": "",
        "sections": [],
        "success": False,
        "error": None
    }

    try:
        doc = Document(io.BytesIO(file_bytes))

        paragraphs_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

        # Extract text from tables (resumes often use tables for layout)
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_texts.append(" | ".join(row_data))

        # Extract from headers/footers
        header_texts = []
        for section in doc.sections:
            header = section.header
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        header_texts.append(para.text.strip())

        # Extract hyperlinks from document relationships (LinkedIn, GitHub URLs etc.)
        hyperlink_urls = []
        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in str(rel.reltype).lower():
                    url = rel.target_ref
                    if url and isinstance(url, str) and url.startswith('http'):
                        hyperlink_urls.append(url)
        except Exception as e:
            logger.warning(f"Could not extract hyperlinks from DOCX: {e}")

        # Combine all text sources
        all_text_parts = header_texts + paragraphs_text + table_texts
        # Append hyperlinks so the entity extractor can find them
        if hyperlink_urls:
            all_text_parts.append("\n".join(hyperlink_urls))
        result["text"] = "\n".join(all_text_parts).strip()
        result["sections"] = paragraphs_text

        if len(result["text"]) < 20:
            result["error"] = "Very little text extracted from DOCX"
            result["success"] = len(result["text"].strip()) > 0
            return result

        result["success"] = True

    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        result["error"] = f"DOCX parsing error: {str(e)}"

    return result
